"""Extraction de texte depuis les documents de consultation (DCE).

PDF via pdfplumber, DOCX via python-docx, dézippage des archives ZIP.

Pas d'OCR implémenté : un PDF dont le texte extrait est trop court (< 50
caractères) est marqué "PDF probablement scanné" plutôt que de produire un
texte vide en silence (règle d'exactitude — ne jamais masquer une limite).
"""
from __future__ import annotations

import logging
import zipfile
from pathlib import Path

import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)

SEUIL_TEXTE_SCANNE = 50

# Documents génériques à exclure du téléchargement / de la synthèse
# (section 6 du cahier des charges).
GENERIC_DOCUMENT_NAMES = (
    "reglement de la consultation", "règlement de la consultation",
    "acte d'engagement", "acte engagement",
    "cgu", "conditions generales", "conditions générales",
    "faq", "guide", "depot de pli", "dépôt de pli",
    "guide utilisateur", "notice",
)


class ExtractionResult:
    """Résultat d'extraction : texte + statut explicite (jamais un texte vide
    silencieux)."""

    def __init__(self, text: str, status: str, source_path: Path):
        self.text = text
        self.status = status  # "ok" | "vide" | "PDF probablement scanné" | "erreur" | "type non supporté"
        self.source_path = source_path

    def __repr__(self) -> str:
        return f"ExtractionResult(status={self.status!r}, len={len(self.text)}, path={self.source_path})"


def extract_text_from_pdf(path: Path) -> ExtractionResult:
    try:
        chunks = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                chunks.append(page.extract_text() or "")
        text = "\n".join(chunks).strip()
    except Exception as exc:  # pdfplumber/pdfminer peuvent lever divers types selon le PDF
        logger.warning("Erreur extraction PDF %s : %s", path, exc)
        return ExtractionResult("", "erreur", path)

    if len(text) < SEUIL_TEXTE_SCANNE:
        return ExtractionResult(text, "PDF probablement scanné", path)
    return ExtractionResult(text, "ok", path)


def extract_text_from_docx(path: Path) -> ExtractionResult:
    try:
        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        text = "\n".join(p for p in paragraphs if p).strip()
    except Exception as exc:
        logger.warning("Erreur extraction DOCX %s : %s", path, exc)
        return ExtractionResult("", "erreur", path)

    if not text:
        return ExtractionResult(text, "vide", path)
    return ExtractionResult(text, "ok", path)


def extract_zip(path: Path, dest_dir: Path) -> list[Path]:
    """Dézippe une archive DCE. Retourne la liste des fichiers extraits.
    Seul le nom de fichier (sans arborescence) est conservé, ce qui élimine
    tout risque de traversée de chemin (path traversal) via un nom de membre
    malveillant."""
    dest_dir.mkdir(parents=True, exist_ok=True)
    extracted: list[Path] = []
    try:
        with zipfile.ZipFile(path) as zf:
            for member in zf.namelist():
                if member.endswith("/"):
                    continue
                safe_name = Path(member).name
                if not safe_name:
                    continue
                target = dest_dir / safe_name
                with zf.open(member) as src, open(target, "wb") as dst:
                    dst.write(src.read())
                extracted.append(target)
    except zipfile.BadZipFile as exc:
        logger.warning("Archive ZIP invalide %s : %s", path, exc)
    return extracted


def is_generic_document(filename: str) -> bool:
    """Détecte les documents génériques à exclure (guides, CGU, FAQ...)."""
    name_normalized = Path(filename).stem.lower().replace("_", " ").replace("-", " ")
    return any(generic in name_normalized for generic in GENERIC_DOCUMENT_NAMES)


def extract_text(path: Path) -> ExtractionResult:
    """Point d'entrée générique : dispatch selon l'extension du fichier."""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_text_from_pdf(path)
    if suffix == ".docx":
        return extract_text_from_docx(path)
    return ExtractionResult("", "type non supporté", path)
