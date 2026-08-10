"""Tests de extractor.py — PDF/DOCX/ZIP, détection de documents génériques."""
from __future__ import annotations

import zipfile
from pathlib import Path

import pytest
from docx import Document

import extractor


def test_extract_text_from_docx_reel(tmp_path: Path):
    docx_path = tmp_path / "cctp.docx"
    doc = Document()
    doc.add_paragraph("Cahier des clauses techniques particulières")
    doc.add_paragraph("Développement d'une application de gestion RH")
    doc.save(docx_path)

    result = extractor.extract_text_from_docx(docx_path)
    assert result.status == "ok"
    assert "gestion RH" in result.text


def test_extract_text_from_docx_vide(tmp_path: Path):
    docx_path = tmp_path / "vide.docx"
    Document().save(docx_path)

    result = extractor.extract_text_from_docx(docx_path)
    assert result.status == "vide"


def test_pdf_probablement_scanne_si_texte_court(tmp_path: Path, monkeypatch):
    """Simule un PDF scanné (image sans couche texte) : pdfplumber renvoie
    un texte quasi vide sur chaque page."""

    class FakePage:
        def extract_text(self):
            return ""

    class FakePdf:
        pages = [FakePage(), FakePage()]

        def __enter__(self):
            return self

        def __exit__(self, *args):
            return False

    monkeypatch.setattr(extractor.pdfplumber, "open", lambda path: FakePdf())

    fake_path = tmp_path / "scan.pdf"
    fake_path.write_bytes(b"%PDF-1.4 fake")
    result = extractor.extract_text_from_pdf(fake_path)
    assert result.status == "PDF probablement scanné"


def test_pdf_texte_normal_extrait(tmp_path: Path, monkeypatch):
    long_text = "Objet du marché : développement logiciel. " * 5  # > 50 caractères

    class FakePage:
        def extract_text(self):
            return long_text

    class FakePdf:
        pages = [FakePage()]

        def __enter__(self):
            return self

        def __exit__(self, *args):
            return False

    monkeypatch.setattr(extractor.pdfplumber, "open", lambda path: FakePdf())

    fake_path = tmp_path / "cctp.pdf"
    fake_path.write_bytes(b"%PDF-1.4 fake")
    result = extractor.extract_text_from_pdf(fake_path)
    assert result.status == "ok"
    assert "développement logiciel" in result.text


def test_extract_zip_extrait_les_fichiers(tmp_path: Path):
    zip_path = tmp_path / "dce.zip"
    with zipfile.ZipFile(zip_path, "w") as zf:
        zf.writestr("CCTP.pdf", b"contenu pdf factice")
        zf.writestr("sous_dossier/RC.docx", b"contenu docx factice")

    dest_dir = tmp_path / "extrait"
    extracted = extractor.extract_zip(zip_path, dest_dir)

    names = sorted(p.name for p in extracted)
    assert names == ["CCTP.pdf", "RC.docx"]
    assert (dest_dir / "CCTP.pdf").read_bytes() == b"contenu pdf factice"


@pytest.mark.parametrize("filename", [
    "Guide_utilisateur_PLACE.pdf",
    "CGU_plateforme.pdf",
    "FAQ.pdf",
    "reglement_de_la_consultation.pdf",
    "depot_de_pli.pdf",
])
def test_is_generic_document_detecte_les_documents_generiques(filename):
    assert extractor.is_generic_document(filename) is True


@pytest.mark.parametrize("filename", [
    "CCTP.pdf",
    "CCAP.pdf",
    "DCE_lot1.pdf",
    "Memoire_technique_type.docx",
])
def test_is_generic_document_ne_detecte_pas_les_documents_utiles(filename):
    assert extractor.is_generic_document(filename) is False


def test_extract_text_type_non_supporte(tmp_path: Path):
    path = tmp_path / "fichier.xyz"
    path.write_text("contenu")
    result = extractor.extract_text(path)
    assert result.status == "type non supporté"
