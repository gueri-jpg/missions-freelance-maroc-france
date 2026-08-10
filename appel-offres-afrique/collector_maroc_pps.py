"""Collecte du "Programme Prévisionnel" (PPS) — Maroc.

Source : https://www.marchespublics.gov.ma/index.php?page=entreprise.ListePPs
Équivalent marocain d'APProch/PPM : documents annonçant les achats prévus
par chaque acheteur public avant publication formelle des avis. **Structure
fondamentalement différente du PPM ivoirien** (`collector_ci_ppm.py`) :

- Le PPM ivoirien est UN document centralisé, structuré en table exploitable
  (une ligne = un projet, colonnes homogènes) publié par la DGMP elle-même.
- Le PPS marocain est un DÉPÔT de ~5 600+ documents PDF individuels,
  téléversés indépendamment par chaque acheteur public (une ligne du site =
  un fichier, pas un projet), sans table structurée exploitable.

Constaté en testant un échantillon réel : ces "PDF" sont d'une hétérogénéité
majeure. Sur 15 documents récents inspectés en détail (05/08/2026) :
- ~1/4 n'étaient même pas des PDF malgré le nom du champ — de vrais
  fichiers `.xlsx`/`.docx`/`.zip` téléversés tels quels par l'acheteur.
  Rejetés à tort auparavant (aucun rapport avec un scan illisible, juste un
  format non géré) — `_extract_text_from_xlsx`/`_extract_text_from_docx`
  corrigent ce point sans dépendance système (openpyxl déjà utilisé par le
  projet, python-docx en pur Python).
- Le reste des PDF sans texte sont de VRAIS SCANS (un seul objet image par
  page). **OCR délibérément écarté après évaluation empirique** (demandé
  explicitement, 05/08/2026) : un échantillon rendu en image et inspecté
  visuellement s'est avéré parfaitement lisible techniquement, mais son
  contenu réel ("Travaux de raccordement électrique souterrain... commune
  de Guisser") confirme que ces scans proviennent presque systématiquement
  de petites entités (communes rurales, directions provinciales) dont la
  commande publique est structurellement BTP/électricité/eau, pas IT — le
  gain attendu d'un OCR ne justifiait pas le coût (dépendance système
  Tesseract, ou coût récurrent par document pour un LLM vision). Ces scans
  restent donc silencieusement écartés (jamais classés par défaut),
  conformément à la règle d'exactitude du projet.

Ce module reste volontairement BORNÉ (`max_documents`, les plus récents
d'abord — la liste est déjà triée par date de publication décroissante côté
serveur) : télécharger et tenter de parser l'intégralité du dépôt (~5 600
documents) représenterait une charge disproportionnée sur le serveur cible
pour un rendement incertain. Chaque document conservé pointe vers son PDF
d'origine (à consulter manuellement) — ce n'est PAS une extraction de projets
individuels comme pour le PPM ivoirien.
"""
from __future__ import annotations

import io
import logging
import re
import time

import openpyxl
import pdfplumber
import requests
from docx import Document
from playwright.sync_api import sync_playwright

import config

logger = logging.getLogger(__name__)


class MarocPpsError(Exception):
    """Erreur de communication ou de structure inattendue pour la source PPS."""


_APOSTROPHE_VARIANTS = str.maketrans({"'": "'", "'": "'", "´": "'", "`": "'"})


def _normalize(text: str) -> str:
    return text.translate(_APOSTROPHE_VARIANTS).lower()


def _find_keyword_snippet(text: str, keywords: list[str], context_chars: int = 150) -> str | None:
    """Retourne un extrait du texte autour de la première occurrence d'un
    mot-clé IT, ou None si aucun mot-clé n'est trouvé. Cet extrait sert
    d'"objet" faute de structure exploitable dans le document source."""
    normalized = _normalize(text)
    for kw in keywords:
        idx = normalized.find(kw)
        if idx != -1:
            start = max(0, idx - context_chars // 2)
            end = min(len(text), idx + len(kw) + context_chars // 2)
            snippet = " ".join(text[start:end].split())
            return f"…{snippet}…" if start > 0 or end < len(text) else snippet
    return None


_HREF_ID_ORG_RE = re.compile(r"id=(\d+)&(?:amp;)?org=([a-zA-Z0-9]+)")

_JS_EXTRACT_PPS_ROWS = r"""
() => {
  const rows = Array.from(document.querySelectorAll('tr')).filter(
    tr => tr.querySelector('td[headers="nomFichier"]')
  );
  return rows.map(row => {
    const q = (sel) => row.querySelector(sel);
    const text = (el) => el ? el.textContent.trim().replace(/\s+/g, ' ') : null;
    const fileLink = q('td[headers="nomFichier"] a');
    return {
      nom_fichier: text(fileLink),
      href_raw: fileLink ? fileLink.getAttribute('href') : null,
      acheteur: text(q('td[headers="PathService"] span')),
      annee: text(q('td[headers="annee"] span')),
      date_publication: text(q('td[headers="datePublication"] span')),
    };
  });
}
"""


def _clean_date(value: str | None) -> str | None:
    if not value:
        return None
    m = re.match(r"(\d{2})/(\d{2})/(\d{4})", value.strip())
    if not m:
        return None
    day, month, year = m.groups()
    return f"{year}-{month}-{day}"


def fetch_document_index(max_documents: int) -> list[dict]:
    """Soumet la recherche du programme prévisionnel (sans filtre — tous
    acheteurs, toutes années) et retourne les `max_documents` entrées les
    plus récentes (la liste est déjà triée par date de publication
    décroissante côté serveur — vérifié en direct)."""
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        context = browser.new_context(
            user_agent=(
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                "(KHTML, like Gecko) Chrome/124.0 Safari/537.36"
            ),
            locale="fr-FR",
        )
        page = context.new_page()
        try:
            page.goto(config.MA_PMMP_PPS_URL, timeout=45000, wait_until="load")
            page.wait_for_timeout(2000)
            submit = page.locator('input[name="ctl0$CONTENU_PAGE$ctl4"]')
            submit.first.click()
            # Le postback PRADO ne déclenche pas toujours un évènement de
            # navigation classique détectable par Playwright — on attend
            # plutôt activement l'apparition des lignes de résultat (jusqu'à
            # ~15s), avec un repli sur un nouveau clic si rien n'apparaît
            # (constaté en direct : un premier essai peut ne rien renvoyer
            # sans erreur explicite).
            row_locator = page.locator('td[headers="nomFichier"]')
            try:
                row_locator.first.wait_for(state="attached", timeout=15000)
            except Exception:
                submit.first.click()
                row_locator.first.wait_for(state="attached", timeout=15000)
            page.wait_for_timeout(1500)
            raw_rows = page.evaluate(_JS_EXTRACT_PPS_ROWS)
        except Exception as exc:
            raise MarocPpsError(f"Erreur de navigation PPS : {exc}") from exc
        finally:
            browser.close()

    return raw_rows[:max_documents]


def _extract_text_from_pdf(content: bytes, max_pages: int) -> str:
    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            texts = [p.extract_text() or "" for p in pdf.pages[:max_pages]]
        return "\n".join(texts)
    except Exception:
        return ""


def _extract_text_from_xlsx(content: bytes) -> str:
    """Best-effort : concatène toutes les valeurs de cellules non vides,
    toutes feuilles confondues. Chaîne vide si le fichier n'est pas un
    xlsx valide (jamais d'exception)."""
    try:
        wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    except Exception:
        return ""
    parts = [
        str(value)
        for ws in wb.worksheets
        for row in ws.iter_rows(values_only=True)
        for value in row
        if value is not None
    ]
    return " ".join(parts)


def _extract_text_from_docx(content: bytes) -> str:
    """Best-effort : paragraphes + cellules de tableaux. Chaîne vide si le
    fichier n'est pas un docx valide (jamais d'exception)."""
    try:
        doc = Document(io.BytesIO(content))
    except Exception:
        return ""
    parts = [p.text for p in doc.paragraphs if p.text]
    parts += [cell.text for table in doc.tables for row in table.rows for cell in row.cells if cell.text]
    return " ".join(parts)


def download_and_extract_text(doc_id: str, org: str, max_pages: int = 6) -> str:
    """Télécharge un document PPS et tente une extraction de texte (bornée
    aux `max_pages` premières pages pour les PDF — certains documents font
    plusieurs dizaines de pages, coût d'extraction disproportionné pour un
    simple filtrage par mot-clé). Gère PDF, xlsx et docx (constaté en
    direct : une partie réelle des documents PPS ne sont pas des PDF malgré
    le nom du champ, cf. docstring du module) — jamais d'exception qui
    interromprait la collecte pour un seul document défaillant. Retourne
    une chaîne vide pour un scan (PDF sans texte, OCR délibérément écarté)
    ou tout autre format non géré (zip générique...)."""
    url = (
        f"{config.MA_PMMP_BASE_URL}/index.php"
        f"?page=commun.PopupListePPsDownloadFile&id={doc_id}&org={org}"
    )
    headers = {"User-Agent": config.USER_AGENT, "Referer": config.MA_PMMP_PPS_URL}
    try:
        resp = requests.get(url, headers=headers, timeout=config.HTTP_TIMEOUT)
    except requests.RequestException as exc:
        logger.warning("PPS %s : erreur réseau (non bloquant) : %s", doc_id, exc)
        return ""
    if not resp.ok:
        return ""

    content = resp.content
    if content.startswith(b"%PDF"):
        return _extract_text_from_pdf(content, max_pages)
    if content[:2] == b"PK":
        # xlsx/docx/zip partagent tous le conteneur ZIP (format OOXML) —
        # on essaie xlsx puis docx, chacun échoue proprement (chaîne vide)
        # si ce n'est pas le bon format ; un zip générique (constaté en
        # direct) retombe sur "" comme avant ce correctif.
        text = _extract_text_from_xlsx(content)
        return text if text else _extract_text_from_docx(content)
    return ""


def collect(keywords: list[str] | None = None, max_documents: int = 30) -> list[dict]:
    """Parcourt les `max_documents` documents PPS les plus récents,
    télécharge et tente d'en extraire le texte, et ne conserve que ceux où
    un mot-clé IT est trouvé. Les documents sans texte extractible (scans,
    PDF invalides) sont silencieusement écartés — jamais classés "hors IT"
    par défaut (règle d'exactitude : l'absence de lecture ne prouve rien)."""
    keywords = keywords or config.MOTS_CLES_IT
    rows = fetch_document_index(max_documents)
    logger.info("PPS Maroc : %d document(s) récents examinés (sur ~5 600+ au total)", len(rows))

    results: list[dict] = []
    n_illisibles = 0
    for i, raw in enumerate(rows):
        href = raw.get("href_raw") or ""
        m = _HREF_ID_ORG_RE.search(href)
        if not m:
            continue
        doc_id, org = m.groups()

        if i > 0:
            # Délai plus long que REQUEST_DELAY_SECONDS (1.5s) : ce portail a
            # montré des signes de ralentissement (multiples "Read timed
            # out") sous accès séquentiel soutenu lors des tests réels — cf.
            # README, Limites de couverture.
            time.sleep(max(config.REQUEST_DELAY_SECONDS, 3.0))
        text = download_and_extract_text(doc_id, org)
        if not text.strip():
            n_illisibles += 1
            continue

        snippet = _find_keyword_snippet(text, keywords)
        if snippet is None:
            continue

        results.append(
            {
                "id": f"MA-PPS-{doc_id}",
                "pays": "Maroc",
                "source": "PMMP (programme prévisionnel)",
                "reference": doc_id,
                "objet": f"[Document non structuré] {snippet}",
                "acheteur": raw.get("acheteur") or "non précisé",
                "ministere": None,
                "bailleur": None,
                "type_marche": None,
                "mode_passation": None,
                "date_publication": _clean_date(raw.get("date_publication")),
                "date_limite": None,
                "devise": "MAD",
                "montant_estime": None,
                "montant_remarque": (
                    "programme prévisionnel — document non structuré (PDF entier à consulter), "
                    "pas un projet individualisé"
                ),
                "url_avis": (
                    f"{config.MA_PMMP_BASE_URL}/index.php"
                    f"?page=commun.PopupListePPsDownloadFile&id={doc_id}&org={org}"
                ),
            }
        )

    logger.info(
        "PPS Maroc : %d document(s) illisibles/scannés ignorés, %d candidat(s) IT retenus",
        n_illisibles, len(results),
    )
    return results
